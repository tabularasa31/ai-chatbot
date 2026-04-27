"""Tests for document upload and parsing API."""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timedelta, timezone
from unittest.mock import patch

import httpx
import pytest
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from sqlalchemy.orm import Session

from backend.auth.service import create_token_for_user
from backend.tenants.service import create_tenant
from backend.documents import embedder as embedder_mod
from backend.documents import http_client as http_client_mod
from backend.documents.constants import KNOWLEDGE_DOCUMENT_CAPACITY
from backend.core.security import hash_password
from backend.models import (
    Document,
    DocumentStatus,
    DocumentType,
    Embedding,
    QuickAnswer,
    SourceSchedule,
    SourceStatus,
    UrlSource,
    UrlSourceRun,
)
from tests.conftest import register_and_verify_user


def _make_minimal_pdf() -> bytes:
    """Create a minimal valid PDF in memory for testing."""
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _fake_embedding_vector() -> list[float]:
    return [0.1] * 1536


def _get_unverified_user_token(db_session: Session, email: str) -> str:
    from backend.models import User

    user = User(
        email=email,
        password_hash=hash_password("SecurePass1!"),
        is_verified=False,
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    token, _ = create_token_for_user(user)
    return token


def test_upload_pdf_success(tenant: TestClient, db_session: Session) -> None:
    """Upload a small PDF, get DocumentResponse back, status=ready."""
    token = register_and_verify_user(tenant, db_session, email="pdf@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "PDF Tenant"},
    )
    pdf_bytes = _make_minimal_pdf()
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
    )
    assert response.status_code == 201
    data = response.json()
    assert "id" in data
    assert data["filename"] == "test.pdf"
    assert data["file_type"] == "pdf"
    assert data["status"] == "ready"
    assert "created_at" in data
    assert "updated_at" in data


def test_upload_markdown_success(tenant: TestClient, db_session: Session) -> None:
    """Upload .md file, status=ready."""
    token = register_and_verify_user(tenant, db_session, email="md@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "MD Tenant"},
    )
    md_content = b"# Test\n\nThis is a test document."
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("test.md", md_content, "text/markdown")},
    )
    assert response.status_code == 201
    data = response.json()
    assert data["filename"] == "test.md"
    assert data["file_type"] == "markdown"
    assert data["status"] == "ready"


def test_upload_swagger_success(tenant: TestClient, db_session: Session) -> None:
    """Upload valid OpenAPI JSON, status=ready."""
    token = register_and_verify_user(tenant, db_session, email="swagger@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Swagger Tenant"},
    )
    swagger_content = b'{"openapi":"3.0.0","info":{"title":"Test API","version":"1.0"},"paths":{"/test":{"get":{"description":"Test endpoint"}}}}'
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("api.json", swagger_content, "application/json")},
    )
    assert response.status_code == 201
    data = response.json()
    assert data["filename"] == "api.json"
    assert data["file_type"] == "swagger"
    assert data["status"] == "ready"


def test_upload_swagger_yaml_success(tenant: TestClient, db_session: Session) -> None:
    """Upload valid OpenAPI YAML, status=ready."""
    token = register_and_verify_user(tenant, db_session, email="swagger-yaml@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Swagger YAML Tenant"},
    )
    swagger_content = b"""
openapi: 3.0.0
info:
  title: YAML Test API
  version: "1.0"
paths:
  /users:
    post:
      summary: Create user
      responses:
        "201":
          description: Created
"""
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("api.yaml", swagger_content, "application/yaml")},
    )
    assert response.status_code == 201
    data = response.json()
    assert data["filename"] == "api.yaml"
    assert data["file_type"] == "swagger"
    assert data["status"] == "ready"


def _make_minimal_docx() -> bytes:
    """Create a minimal valid .docx file in memory for testing."""
    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_paragraph("Hello from docx test document.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_upload_docx_success(tenant: TestClient, db_session: Session) -> None:
    """Upload .docx file, status=ready."""
    token = register_and_verify_user(tenant, db_session, email="docx@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "DOCX Tenant"},
    )
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("report.docx", _make_minimal_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201
    data = response.json()
    assert data["filename"] == "report.docx"
    assert data["file_type"] == "docx"
    assert data["status"] == "ready"


def test_upload_txt_success(tenant: TestClient, db_session: Session) -> None:
    """Upload .txt file, status=ready."""
    token = register_and_verify_user(tenant, db_session, email="txt@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "TXT Tenant"},
    )
    txt_content = b"This is a plain text document.\n\nIt has multiple paragraphs."
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("notes.txt", txt_content, "text/plain")},
    )
    assert response.status_code == 201
    data = response.json()
    assert data["filename"] == "notes.txt"
    assert data["file_type"] == "plaintext"
    assert data["status"] == "ready"


def test_upload_unsupported_type(tenant: TestClient, db_session: Session) -> None:
    """Upload .exe → 400."""
    token = register_and_verify_user(tenant, db_session, email="exe@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Exe Tenant"},
    )
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("virus.exe", b"MZ", "application/octet-stream")},
    )
    assert response.status_code == 400
    assert "unsupported" in response.json()["detail"].lower()


def test_upload_too_large(tenant: TestClient, db_session: Session) -> None:
    """Upload >50MB file → 400."""
    token = register_and_verify_user(tenant, db_session, email="large@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Large Tenant"},
    )
    # 51 MB
    large_content = b"x" * (51 * 1024 * 1024)
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("huge.pdf", large_content, "application/pdf")},
    )
    assert response.status_code == 400
    assert "too large" in response.json()["detail"].lower()


def test_upload_no_client(tenant: TestClient, db_session: Session) -> None:
    """Upload without creating tenant first → 404."""
    token = register_and_verify_user(tenant, db_session, email="noclient@example.com")
    md_content = b"# Test"
    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("test.md", md_content, "text/markdown")},
    )
    assert response.status_code == 404
    assert "tenant" in response.json()["detail"].lower()


def test_upload_document_limit_is_shared_capacity(tenant: TestClient, db_session: Session) -> None:
    """The tenant-wide document capacity should fill up and reject the next upload."""
    token = register_and_verify_user(tenant, db_session, email="limit100@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Limit Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    for index in range(KNOWLEDGE_DOCUMENT_CAPACITY):
        db_session.add(
            Document(
                tenant_id=tenant_id,
                filename=f"existing-{index}.md",
                file_type=DocumentType.markdown,
                status=DocumentStatus.ready,
                parsed_text=f"doc {index}",
            )
        )
    db_session.commit()

    response = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("overflow.md", b"# Overflow", "text/markdown")},
    )

    assert response.status_code == 400
    assert response.json()["detail"] == f"Document limit reached (max {KNOWLEDGE_DOCUMENT_CAPACITY})"


def test_upload_unauthenticated(tenant: TestClient) -> None:
    """No JWT → 401."""
    md_content = b"# Test"
    response = tenant.post(
        "/documents",
        files={"file": ("test.md", md_content, "text/markdown")},
    )
    assert response.status_code == 401


def test_list_documents_empty(tenant: TestClient, db_session: Session) -> None:
    """Get documents when none uploaded → empty list."""
    token = register_and_verify_user(tenant, db_session, email="empty@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Empty Tenant"},
    )
    response = tenant.get(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    assert "documents" in data
    assert data["documents"] == []


def test_list_knowledge_sources_requires_client(tenant: TestClient, db_session: Session) -> None:
    """Knowledge sources should match other document routes and 404 without a tenant."""
    token = register_and_verify_user(tenant, db_session, email="sources-noclient@example.com")

    response = tenant.get(
        "/documents/sources",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 404
    assert response.json()["detail"] == "Tenant not found"


def test_list_documents(tenant: TestClient, db_session: Session) -> None:
    """Upload 2 docs, get list → 2 items."""
    token = register_and_verify_user(tenant, db_session, email="list@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "List Tenant"},
    )
    md1 = b"# Doc 1"
    md2 = b"# Doc 2"
    tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("doc1.md", md1, "text/markdown")},
    )
    tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("doc2.md", md2, "text/markdown")},
    )
    response = tenant.get(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    assert len(data["documents"]) == 2
    filenames = {d["filename"] for d in data["documents"]}
    assert filenames == {"doc1.md", "doc2.md"}


def test_get_document_detail(tenant: TestClient, db_session: Session) -> None:
    """Get single document, verify parsed_text preview."""
    token = register_and_verify_user(tenant, db_session, email="detail@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Detail Tenant"},
    )
    md_content = b"# Test\n\nThis is a test document with some content."
    upload_resp = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("detail.md", md_content, "text/markdown")},
    )
    doc_id = upload_resp.json()["id"]
    response = tenant.get(
        f"/documents/{doc_id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["id"] == str(doc_id)
    assert data["filename"] == "detail.md"
    assert "parsed_text" in data
    assert "Test" in (data["parsed_text"] or "")
    assert "test document" in (data["parsed_text"] or "")


def test_get_document_wrong_user(tenant: TestClient, db_session: Session) -> None:
    """User B tries to get user A's document → 404."""
    token_a = register_and_verify_user(tenant, db_session, email="userA@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token_a}"},
        json={"name": "User A Tenant"},
    )
    md_content = b"# Secret"
    upload_resp = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token_a}"},
        files={"file": ("secret.md", md_content, "text/markdown")},
    )
    doc_id = upload_resp.json()["id"]

    token_b = register_and_verify_user(tenant, db_session, email="userB@example.com")

    response = tenant.get(
        f"/documents/{doc_id}",
        headers={"Authorization": f"Bearer {token_b}"},
    )
    assert response.status_code == 404


def test_get_url_source_detail_returns_latest_five_runs(
    tenant: TestClient, db_session: Session
) -> None:
    """Source detail should return the five newest runs ordered in SQL."""
    token = register_and_verify_user(tenant, db_session, email="source-runs@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Source Runs Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.weekly,
        pages_indexed=0,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.flush()

    base_time = datetime.now(timezone.utc)
    expected_statuses: list[str] = []
    for index in range(7):
        run = UrlSourceRun(
            source_id=source.id,
            status=f"run-{index}",
            pages_indexed=index,
            failed_urls=[],
            created_at=base_time + timedelta(minutes=index),
        )
        db_session.add(run)
        if index >= 2:
            expected_statuses.insert(0, run.status)
    db_session.commit()

    response = tenant.get(
        f"/documents/sources/{source.id}",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    recent_runs = response.json()["recent_runs"]
    assert len(recent_runs) == 5
    assert [run["status"] for run in recent_runs] == expected_statuses


@patch("backend.auth.routes.send_email")
def test_delete_url_source_requires_verified_user(
    mock_send_email, tenant: TestClient, db_session: Session
) -> None:
    token = _get_unverified_user_token(db_session, "unverified-source-delete@example.com")
    from backend.models import User

    user = db_session.query(User).filter(User.email == "unverified-source-delete@example.com").first()
    assert user is not None
    owner_client = create_tenant(user.id, "Unverified Tenant", db_session)

    source = UrlSource(
        tenant_id=owner_client.id,
        name="Docs",
        url="https://docs.example.com/",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.weekly,
        pages_indexed=0,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.commit()

    response = tenant.delete(
        f"/documents/sources/{source.id}",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 403
    assert response.json()["detail"] == "Email not verified."


def test_get_url_source_detail_includes_quick_answers(
    tenant: TestClient, db_session: Session
) -> None:
    token = register_and_verify_user(tenant, db_session, email="source-qa@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Source QA Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.weekly,
        pages_indexed=1,
        chunks_created=2,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.flush()
    db_session.add(
        QuickAnswer(
            tenant_id=tenant_id,
            source_id=source.id,
            key="documentation_url",
            value="https://docs.example.com/",
            source_url="https://docs.example.com/",
            metadata_json={"method": "source_url"},
        )
    )
    db_session.commit()

    response = tenant.get(
        f"/documents/sources/{source.id}",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["quick_answers"] == [
        {
            "key": "documentation_url",
            "value": "https://docs.example.com/",
            "source_url": "https://docs.example.com/",
            "detected_at": payload["quick_answers"][0]["detected_at"],
        }
    ]


def test_create_url_source_rejects_duplicate_normalized_domain(
    tenant: TestClient, db_session: Session, monkeypatch: pytest.MonkeyPatch
) -> None:
    token = register_and_verify_user(tenant, db_session, email="source-dup@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Source Dup Tenant"},
    )

    monkeypatch.setattr(
        "backend.documents.http_client._fetch_reachable_page",
        lambda url, timeout_seconds: ("<html></html>", "Docs"),
    )
    monkeypatch.setattr("backend.documents.http_client._validate_public_hostname", lambda hostname: None)
    monkeypatch.setattr("backend.documents.sitemap._load_robots_warning", lambda url: None)
    monkeypatch.setattr(
        "backend.documents.url_service._discover_urls",
        lambda root_url, exclusions, page_cap: [root_url],
    )

    first_response = tenant.post(
        "/documents/sources/url",
        headers={"Authorization": f"Bearer {token}"},
        json={"url": "https://docs.example.com/start", "schedule": "manual"},
    )
    second_response = tenant.post(
        "/documents/sources/url",
        headers={"Authorization": f"Bearer {token}"},
        json={"url": "https://docs.example.com/another", "schedule": "manual"},
    )

    assert first_response.status_code == 201
    assert second_response.status_code == 409
    assert "already have a source from this domain" in second_response.json()["detail"].lower()


def test_refresh_url_source_returns_429_inside_cooldown(
    tenant: TestClient, db_session: Session, monkeypatch: pytest.MonkeyPatch
) -> None:
    token = register_and_verify_user(tenant, db_session, email="source-refresh@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Source Refresh Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.weekly,
        last_refresh_requested_at=datetime.now(timezone.utc),
        pages_indexed=0,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.commit()

    monkeypatch.setattr("backend.documents.url_service._utcnow", lambda: datetime.now(timezone.utc))

    response = tenant.post(
        f"/documents/sources/{source.id}/refresh",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 429
    assert "refresh available in" in response.json()["detail"].lower()


def test_url_source_crawl_uses_remaining_shared_capacity(
    tenant: TestClient, db_session: Session, monkeypatch: pytest.MonkeyPatch
) -> None:
    """URL pages should consume the same shared capacity pool as uploaded files."""
    from backend.documents import url_service

    monkeypatch.setattr(url_service, "SessionLocal", lambda: db_session)
    monkeypatch.setattr(db_session, "close", lambda: None)

    token = register_and_verify_user(tenant, db_session, email="shared-capacity@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Shared Capacity Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    prefill_count = KNOWLEDGE_DOCUMENT_CAPACITY - 40
    for index in range(prefill_count):
        db_session.add(
            Document(
                tenant_id=tenant_id,
                filename=f"file-{index}.md",
                file_type=DocumentType.markdown,
                status=DocumentStatus.ready,
                parsed_text=f"file {index}",
            )
        )

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/",
        normalized_domain="docs.example.com",
        status=SourceStatus.queued,
        crawl_schedule=SourceSchedule.manual,
        pages_indexed=0,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.commit()

    discovered_urls = [f"https://docs.example.com/page-{index}" for index in range(60)]
    monkeypatch.setattr(url_service, "_discover_urls", lambda *_args, **_kwargs: discovered_urls)
    monkeypatch.setattr(http_client_mod, "_fetch_page_html", lambda url: f"<html>{url}</html>")
    monkeypatch.setattr(
        embedder_mod,
        "_extract_page",
        lambda url, html: type("Page", (), {
            "url": url,
            "title": url.rsplit("/", 1)[-1],
            "text": html,
            "chunks": [{"chunk_text": html, "chunk_index": 0, "section_title": None, "token_count": 1, "content_hash": url, "raw_text": html}],
        })(),
    )
    monkeypatch.setattr(embedder_mod, "_embed_chunks", lambda chunks, api_key: [_fake_embedding_vector() for _ in chunks])

    url_service.crawl_url_source(source.id, api_key="test-key")
    db_session.expire_all()

    refreshed_source = db_session.query(UrlSource).filter(UrlSource.id == source.id).first()
    source_docs = db_session.query(Document).filter(Document.source_id == source.id).all()

    assert refreshed_source is not None
    assert len(source_docs) == 40
    assert refreshed_source.pages_indexed == 40
    assert refreshed_source.warning_message is not None
    assert "Knowledge capacity reached" in refreshed_source.warning_message
    assert db_session.query(Document).filter(Document.tenant_id == tenant_id).count() == KNOWLEDGE_DOCUMENT_CAPACITY


def test_url_source_refresh_updates_existing_pages_without_exceeding_shared_capacity(
    tenant: TestClient, db_session: Session, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Refresh should update existing source pages and only add new pages while capacity remains."""
    from backend.documents import url_service

    monkeypatch.setattr(url_service, "SessionLocal", lambda: db_session)
    monkeypatch.setattr(db_session, "close", lambda: None)

    token = register_and_verify_user(tenant, db_session, email="refresh-capacity@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Refresh Capacity Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.manual,
        pages_indexed=60,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.flush()

    for index in range(KNOWLEDGE_DOCUMENT_CAPACITY - 60):
        db_session.add(
            Document(
                tenant_id=tenant_id,
                filename=f"other-{index}.md",
                file_type=DocumentType.markdown,
                status=DocumentStatus.ready,
                parsed_text=f"other {index}",
            )
        )
    for index in range(60):
        db_session.add(
            Document(
                tenant_id=tenant_id,
                source_id=source.id,
                source_url=f"https://docs.example.com/page-{index}",
                filename=f"page-{index}",
                file_type=DocumentType.url,
                status=DocumentStatus.ready,
                parsed_text=f"old {index}",
            )
        )
    db_session.commit()

    discovered_urls = [f"https://docs.example.com/page-{index}" for index in range(70)]
    monkeypatch.setattr(url_service, "_discover_urls", lambda *_args, **_kwargs: discovered_urls)
    monkeypatch.setattr(http_client_mod, "_fetch_page_html", lambda url: f"<html>{url}</html>")
    monkeypatch.setattr(
        embedder_mod,
        "_extract_page",
        lambda url, html: type("Page", (), {
            "url": url,
            "title": url.rsplit("/", 1)[-1],
            "text": f"updated {url}",
            "chunks": [{"chunk_text": html, "chunk_index": 0, "section_title": None, "token_count": 1, "content_hash": url, "raw_text": html}],
        })(),
    )
    monkeypatch.setattr(embedder_mod, "_embed_chunks", lambda chunks, api_key: [_fake_embedding_vector() for _ in chunks])

    url_service.crawl_url_source(source.id, api_key="test-key")
    db_session.expire_all()

    refreshed_source = db_session.query(UrlSource).filter(UrlSource.id == source.id).first()
    source_docs = db_session.query(Document).filter(Document.source_id == source.id).all()

    assert refreshed_source is not None
    assert len(source_docs) == 60
    assert refreshed_source.pages_indexed == 60
    assert refreshed_source.warning_message is not None
    assert "Knowledge capacity reached" in refreshed_source.warning_message
    assert db_session.query(Document).filter(Document.tenant_id == tenant_id).count() == KNOWLEDGE_DOCUMENT_CAPACITY


def test_delete_document_success(tenant: TestClient, db_session: Session) -> None:
    """Delete document → 204, verify gone."""
    token = register_and_verify_user(tenant, db_session, email="del@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Del Tenant"},
    )
    md_content = b"# To Delete"
    upload_resp = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("todel.md", md_content, "text/markdown")},
    )
    doc_id = upload_resp.json()["id"]

    response = tenant.delete(
        f"/documents/{doc_id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 204

    get_resp = tenant.get(
        f"/documents/{doc_id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert get_resp.status_code == 404


def test_delete_document_wrong_user(tenant: TestClient, db_session: Session) -> None:
    """User B tries to delete user A's document → 404."""
    token_a = register_and_verify_user(tenant, db_session, email="delA@example.com")
    tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token_a}"},
        json={"name": "User A Tenant"},
    )
    md_content = b"# Protected"
    upload_resp = tenant.post(
        "/documents",
        headers={"Authorization": f"Bearer {token_a}"},
        files={"file": ("protected.md", md_content, "text/markdown")},
    )
    doc_id = upload_resp.json()["id"]

    token_b = register_and_verify_user(tenant, db_session, email="delB@example.com")

    response = tenant.delete(
        f"/documents/{doc_id}",
        headers={"Authorization": f"Bearer {token_b}"},
    )
    assert response.status_code == 404


def test_delete_source_page_success_persists_exclusion(tenant: TestClient, db_session: Session) -> None:
    token = register_and_verify_user(tenant, db_session, email="delete-source-page@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Delete Source Page Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/start",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.manual,
        pages_found=1,
        pages_indexed=1,
        chunks_created=1,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.flush()

    doc = Document(
        tenant_id=tenant_id,
        source_id=source.id,
        filename="Getting Started",
        file_type=DocumentType.url,
        status=DocumentStatus.ready,
        parsed_text="hello",
        source_url="https://docs.example.com/start",
    )
    db_session.add(doc)
    db_session.flush()
    db_session.add(
        Embedding(
            document_id=doc.id,
            chunk_text="hello",
            vector=_fake_embedding_vector(),
            metadata_json={},
        )
    )
    db_session.commit()

    response = tenant.delete(
        f"/documents/sources/{source.id}/pages/{doc.id}",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 204
    assert db_session.query(Document).filter(Document.id == doc.id).first() is None
    assert db_session.query(Embedding).filter(Embedding.document_id == doc.id).count() == 0

    refreshed_source = db_session.query(UrlSource).filter(UrlSource.id == source.id).first()
    assert refreshed_source is not None
    assert refreshed_source.pages_indexed == 0
    assert refreshed_source.chunks_created == 0
    assert refreshed_source.metadata_json["manually_excluded_page_urls"] == ["https://docs.example.com/start"]


def test_delete_source_page_rejects_document_from_another_source(tenant: TestClient, db_session: Session) -> None:
    token = register_and_verify_user(tenant, db_session, email="delete-source-page-wrong-source@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Delete Source Page Wrong Source Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    first_source = UrlSource(
        tenant_id=tenant_id,
        name="Docs A",
        url="https://docs.example.com/a",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.manual,
        pages_indexed=1,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    second_source = UrlSource(
        tenant_id=tenant_id,
        name="Docs B",
        url="https://docs.example.com/b",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.manual,
        pages_indexed=0,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add_all([first_source, second_source])
    db_session.flush()

    doc = Document(
        tenant_id=tenant_id,
        source_id=first_source.id,
        filename="Page A",
        file_type=DocumentType.url,
        status=DocumentStatus.ready,
        parsed_text="A",
        source_url="https://docs.example.com/a",
    )
    db_session.add(doc)
    db_session.commit()

    response = tenant.delete(
        f"/documents/sources/{second_source.id}/pages/{doc.id}",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 404
    assert db_session.query(Document).filter(Document.id == doc.id).first() is not None


def test_manually_deleted_source_page_is_not_recreated_on_refresh(
    tenant: TestClient, db_session: Session, monkeypatch: pytest.MonkeyPatch
) -> None:
    from backend.documents import url_service

    monkeypatch.setattr(url_service, "SessionLocal", lambda: db_session)
    monkeypatch.setattr(db_session, "close", lambda: None)

    token = register_and_verify_user(tenant, db_session, email="delete-source-page-refresh@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "Delete Source Page Refresh Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="Docs",
        url="https://docs.example.com/start",
        normalized_domain="docs.example.com",
        status=SourceStatus.ready,
        crawl_schedule=SourceSchedule.manual,
        pages_found=1,
        pages_indexed=1,
        chunks_created=1,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.flush()

    doc = Document(
        tenant_id=tenant_id,
        source_id=source.id,
        filename="Getting Started",
        file_type=DocumentType.url,
        status=DocumentStatus.ready,
        parsed_text="hello",
        source_url="https://docs.example.com/start",
    )
    db_session.add(doc)
    db_session.flush()
    db_session.add(
        Embedding(
            document_id=doc.id,
            chunk_text="hello",
            vector=_fake_embedding_vector(),
            metadata_json={},
        )
    )
    db_session.commit()

    delete_response = tenant.delete(
        f"/documents/sources/{source.id}/pages/{doc.id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert delete_response.status_code == 204

    monkeypatch.setattr(url_service, "_discover_urls", lambda root_url, exclusions, page_cap: ["https://docs.example.com/start"])
    monkeypatch.setattr(http_client_mod, "_fetch_page_html", lambda url: "<html><body>start</body></html>")
    monkeypatch.setattr(
        embedder_mod,
        "_extract_page",
        lambda url, html: type(
            "Page",
            (),
            {
                "url": url,
                "title": "Start",
                "text": "start",
                "chunks": [
                    {
                        "chunk_text": "start",
                        "chunk_index": 0,
                        "section_title": None,
                        "token_count": 1,
                        "content_hash": url,
                        "raw_text": html,
                    }
                ],
            },
        )(),
    )
    monkeypatch.setattr(embedder_mod, "_embed_chunks", lambda chunks, api_key: [_fake_embedding_vector() for _ in chunks])

    url_service.crawl_url_source(source.id, api_key="test-key")
    db_session.expire_all()

    refreshed_source = db_session.query(UrlSource).filter(UrlSource.id == source.id).first()
    source_docs = db_session.query(Document).filter(Document.source_id == source.id).all()

    assert refreshed_source is not None
    assert refreshed_source.pages_indexed == 0
    assert refreshed_source.pages_found == 0
    assert source_docs == []


def test_crawl_url_source_detects_openapi_yaml_and_indexes_as_swagger(
    tenant: TestClient, db_session: Session, monkeypatch: pytest.MonkeyPatch
) -> None:
    from backend.documents import url_service

    monkeypatch.setattr(url_service, "SessionLocal", lambda: db_session)
    monkeypatch.setattr(db_session, "close", lambda: None)

    token = register_and_verify_user(tenant, db_session, email="openapi-url@example.com")
    create_client_response = tenant.post(
        "/tenants",
        headers={"Authorization": f"Bearer {token}"},
        json={"name": "OpenAPI URL Tenant"},
    )
    tenant_id = uuid.UUID(create_client_response.json()["id"])

    source = UrlSource(
        tenant_id=tenant_id,
        name="API spec",
        url="https://docs.example.com/openapi.yaml",
        normalized_domain="docs.example.com",
        status=SourceStatus.queued,
        crawl_schedule=SourceSchedule.manual,
        pages_indexed=0,
        chunks_created=0,
        tokens_used=0,
        metadata_json={},
    )
    db_session.add(source)
    db_session.commit()

    monkeypatch.setattr(url_service, "_discover_urls", lambda *_args, **_kwargs: [source.url])
    monkeypatch.setattr(http_client_mod, "_validate_public_hostname", lambda hostname: None)

    yaml_spec = """
openapi: 3.0.0
info:
  title: URL API
  version: "1.0"
paths:
  /users:
    get:
      summary: List users
      operationId: listUsers
      responses:
        "200":
          description: OK
  /users/{userId}:
    get:
      summary: Get user
      parameters:
        - in: path
          name: userId
          required: true
          schema:
            type: string
      responses:
        "200":
          description: OK
"""

    monkeypatch.setattr(
        http_client_mod,
        "_http_client",
        lambda timeout_seconds: httpx.Client(
            transport=httpx.MockTransport(
                lambda request: httpx.Response(
                    200,
                    headers={"content-type": "application/yaml"},
                    text=yaml_spec,
                    request=request,
                )
            ),
            timeout=timeout_seconds,
            follow_redirects=False,
            trust_env=False,
        ),
    )
    monkeypatch.setattr(embedder_mod, "_embed_chunks", lambda chunks, api_key: [_fake_embedding_vector() for _ in chunks])

    url_service.crawl_url_source(source.id, api_key="test-key")
    db_session.expire_all()

    refreshed_source = db_session.query(UrlSource).filter(UrlSource.id == source.id).first()
    doc = db_session.query(Document).filter(Document.source_id == source.id).first()
    embeddings = (
        db_session.query(Embedding)
        .filter(Embedding.document_id == doc.id)
        .order_by(Embedding.created_at.asc())
        .all()
    )

    assert refreshed_source is not None
    assert refreshed_source.status == SourceStatus.ready
    assert refreshed_source.metadata_json["platform"] == "openapi"
    assert doc is not None
    assert doc.file_type == DocumentType.swagger
    assert "Endpoint: GET /users" in (doc.parsed_text or "")
    assert len(embeddings) == 2
    assert embeddings[0].metadata_json["type"] == "api_endpoint"
    assert embeddings[0].metadata_json["source_kind"] == "url"
    assert embeddings[0].metadata_json["path"] == "/users"
